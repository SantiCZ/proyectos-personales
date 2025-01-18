from docx import Document
import re
import pywhatkit as kit
import time

# Cargar el documento
file_path = "MENSAJES SANTI pilar.docx"
doc = Document(file_path)

# Estimar la posición de la página 985 (cada página ~35 párrafos)
start_paragraph = min(985 * 35, len(doc.paragraphs))

# Expresión regular para detectar números telefónicos
phone_pattern = re.compile(r'\b\d{3,4}[-.\s]?\d{6,7}\b')

# Diccionario para almacenar mensajes y números
messages = []
current_message = ""
current_numbers = []

# Extraer mensajes y números desde la página 985
def extract_messages():
    global current_message, current_numbers
    for para in doc.paragraphs[start_paragraph:]:
        matches = phone_pattern.findall(para.text)
        if matches:
            if current_message:
                messages.append((current_message.strip(), current_numbers))
            current_message = ""
            current_numbers = matches
        else:
            current_message += para.text + "\n"
    if current_message:
        messages.append((current_message.strip(), current_numbers))

# Función para enviar mensajes con WhatsApp Web
def send_messages_whatsapp():
    for message, numbers in messages:
        for number in numbers:
            full_number = f"+54{number}"  # Asegurar formato internacional si es Argentina
            kit.sendwhatmsg_instantly(full_number, message, wait_time=15)
            print(f"Mensaje enviado a {full_number}")
            time.sleep(5)  # Esperar unos segundos entre mensajes para evitar bloqueos

# Ejecutar extracción
extract_messages()

# Enviar mensajes usando WhatsApp Web
send_messages_whatsapp()

# Para ver los datos extraídos
for msg, nums in messages:
    print(f"Mensaje: {msg[:50]}...")
    print(f"Numeros: {', '.join(nums)}")
    print("-" * 50)
